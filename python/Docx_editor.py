import sys
import json
import os
import shutil
from docx import Document
from docx.oxml.ns import qn

# Writes edited paragraph text back into an existing .docx file.
#
# Scope / honesty about limitations:
# - This performs a text-level round-trip: it matches edited paragraphs
#   back to the same non-empty paragraphs docx_processor.py extracted,
#   in order, and replaces their text.
# - Run-level formatting (bold/italic/font/color) of the FIRST run in
#   each paragraph is preserved; any additional runs in that paragraph
#   are cleared to avoid duplicated/garbled text. Tables, images,
#   headers/footers, and page layout are left completely untouched.
# - A ".bak" backup of the original file is created (once) before the
#   first save, so edits are always reversible.


def save_docx(file_path, paragraphs):
    if not os.path.isfile(file_path):
        raise FileNotFoundError(f"DOCX file not found: {file_path}")

    if not isinstance(paragraphs, list):
        raise ValueError("paragraphs must be a list of strings")

    backup_path = file_path + ".bak"
    if not os.path.exists(backup_path):
        shutil.copy2(file_path, backup_path)

    document = Document(file_path)

    updated_count = 0
    para_index = 0
    skipped_hyperlink_paragraphs = []

    for paragraph in document.paragraphs:
        # Detect and skip paragraphs containing hyperlinks to avoid data corruption.
        # Paragraphs with <w:hyperlink> elements have empty .runs because the hyperlink runs are nested under the hyperlink element.
        # Therefore we skip editing them and record their indices.
        original_text = paragraph.text.strip()

        if not original_text:
            continue

        if para_index >= len(paragraphs):
            break

        new_text = paragraphs[para_index]

        # Check for hyperlinks in the underlying XML.
        if paragraph._p.findall(qn('w:hyperlink')):
            # Skip editing this paragraph but still count it as processed.
            skipped_hyperlink_paragraphs.append(para_index)
        else:
            # Compare stripped original text to avoid false positives due to whitespace differences.
            if isinstance(new_text, str) and new_text != original_text:
                if paragraph.runs:
                    paragraph.runs[0].text = new_text
                    for extra_run in paragraph.runs[1:]:
                        extra_run.text = ""
                else:
                    paragraph.add_run(new_text)
                updated_count += 1

        para_index += 1

    document.save(file_path)

    return {
        "success": True,
        "filename": os.path.basename(file_path),
        "paragraphs_matched": para_index,
        "paragraphs_changed": updated_count,
        "backup": os.path.basename(backup_path),
        "skipped_hyperlink_paragraphs": skipped_hyperlink_paragraphs,
    }


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print(json.dumps({
            "success": False,
            "error": "Usage: docx_editor.py <docx_path> <payload_json_path>"
        }))
        sys.exit(1)

    docx_path = sys.argv[1]
    payload_path = sys.argv[2]

    try:
        with open(payload_path, "r", encoding="utf-8") as payload_file:
            payload = json.load(payload_file)

        result = save_docx(docx_path, payload.get("paragraphs", []))
        print(json.dumps(result))

    except Exception as error:
        print(json.dumps({
            "success": False,
            "error": str(error)
        }))
        sys.exit(1)